#!/usr/bin/env python3
"""Generate the ATLAS User Guide as a .docx file."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path

doc = Document()

# ── Styles ────────────────────────────────────────────────────────────────

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

# Helper
def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("ATLAS Platform")
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("User Guide")
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x4A, 0x6F, 0xA5)

doc.add_paragraph()
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run("AI-Assisted Criminal Justice Document Platform\nFIR Ingestion · Chargesheet Review · Legal Validation · Evidence Gap Detection")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()
ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
ver.add_run("Version 2.0 — April 2026").font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Getting Started",
    "    1.1 System Requirements",
    "    1.2 Starting the Platform",
    "    1.3 Default User Accounts",
    "    1.4 Logging In",
    "2. Dashboard",
    "3. FIR Module",
    "    3.1 Uploading a FIR PDF",
    "    3.2 Browsing FIRs",
    "    3.3 NLP Classification & Mismatch Detection",
    "4. Chargesheet Module",
    "    4.1 Uploading a Chargesheet PDF",
    "    4.2 Browsing Chargesheets",
    "    4.3 Reviewing a Chargesheet",
    "5. Legal Validation",
    "    5.1 Validation Rules",
    "    5.2 Section Lookup",
    "6. Evidence Gap Detection",
    "    6.1 Rule-Based Detection (Tier 1)",
    "    6.2 ML Pattern Detection (Tier 2)",
    "    6.3 Evidence Taxonomy",
    "7. Review Workflow",
    "    7.1 Starting a Review",
    "    7.2 Acting on Recommendations",
    "    7.3 Completing a Review",
    "8. Audit Trail",
    "    8.1 Viewing the Audit Log",
    "    8.2 Verifying Chain Integrity",
    "    8.3 Exporting for Court",
    "9. Roles & Permissions",
    "10. Service URLs & Monitoring",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if not item.startswith("    "):
        p.runs[0].bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 1. GETTING STARTED
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading("1. Getting Started", level=1)

doc.add_heading("1.1 System Requirements", level=2)
doc.add_paragraph("ATLAS runs as a set of Docker containers. You need:")
reqs = [
    "Docker Desktop (v24+) with at least 8 GB RAM allocated",
    "A modern web browser (Chrome, Firefox, or Edge)",
    "Network access to localhost ports 3000, 8000",
]
for r in reqs:
    doc.add_paragraph(r, style="List Bullet")

doc.add_heading("1.2 Starting the Platform", level=2)
doc.add_paragraph(
    "Open a terminal in the project root directory and run:"
)
doc.add_paragraph("docker compose up -d", style="Intense Quote")
doc.add_paragraph(
    "This starts all services: the backend API, frontend web application, "
    "PostgreSQL database, Redis cache, MongoDB, Prometheus, Grafana, Label Studio, and MLflow. "
    "Wait approximately 30 seconds for all services to become healthy."
)
doc.add_paragraph(
    "To verify all services are running:"
)
doc.add_paragraph("docker compose ps", style="Intense Quote")
doc.add_paragraph(
    'All services should show status "Up" with the backend showing "(healthy)".'
)

doc.add_heading("1.3 Default User Accounts", level=2)
doc.add_paragraph(
    "The system is seeded with three user accounts for testing. "
    "All accounts share the same default password."
)
add_table(
    ["Username", "Password", "Role", "District", "Station"],
    [
        ["admin", "atlas2025", "ADMIN", "Ahmedabad", "—"],
        ["io_sanand", "atlas2025", "IO", "Ahmedabad", "Sanand"],
        ["sho_sanand", "atlas2025", "SHO", "Ahmedabad", "Sanand"],
    ],
)

doc.add_heading("1.4 Logging In", level=2)
steps = [
    "Open your browser and navigate to http://localhost:3000.",
    "You will see the ATLAS login page.",
    'Enter your username and password, then click "Sign In".',
    "On success, you are redirected to the Dashboard.",
    "Your name, username, and role badge appear at the bottom of the left sidebar.",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 2. DASHBOARD
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading("2. Dashboard", level=1)
doc.add_paragraph(
    "The Dashboard is the landing page after login. It displays six real-time metrics "
    "in a card grid, each showing a count and subtitle."
)
add_table(
    ["Card", "Description", "Source"],
    [
        ["Total FIRs", "Total number of FIR records ingested (all time)", "firs table"],
        ["Pending Review", "FIRs with status 'pending' awaiting IO action", "firs table"],
        ["Districts", "Number of distinct districts with FIR data", "firs table"],
        ["Avg Completeness", "Average extraction completeness score (0–100%)", "firs table"],
        ["Ingested Today", "Number of FIR PDFs processed today", "firs table"],
        ["Chargesheets", "Total number of chargesheet documents ingested", "chargesheets table"],
    ],
)
doc.add_paragraph(
    "Below the metrics grid, an amber alert banner confirms that Section Mismatch Detection "
    "is active. FIRs where the NLP narrative analysis contradicts the registered IPC/BNS sections "
    "are automatically flagged as 'review_needed'."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 3. FIR MODULE
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading("3. FIR Module", level=1)
doc.add_paragraph(
    'Navigate to "FIR Review" in the left sidebar. This module handles FIR PDF ingestion, '
    "OCR extraction, NLP classification, and browsing."
)

doc.add_heading("3.1 Uploading a FIR PDF", level=2)
steps = [
    'At the top of the page, you will see a dashed upload zone labelled "Drag & drop a FIR PDF here".',
    "Drag a PDF file onto the zone, or click it to open a file browser.",
    "Only .pdf files are accepted. Non-PDF files will show an error.",
    "While processing, a spinner and the message 'Processing PDF...' appear.",
    "The system extracts text using pdfplumber (for digital PDFs) or OCR via Tesseract (for scanned PDFs) with English and Gujarati language support.",
    "After processing, the extracted data appears in a result card showing: FIR Number, District, Police Station, Sections, Complainant Name, and Completeness percentage.",
    "The narrative (full extracted text) is shown in a separate scrollable card below.",
    "The system automatically classifies the FIR using a 4-tier cascade: section mapping → fine-tuned MuRIL model → zero-shot NLI → keyword heuristics.",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_paragraph(
    "Required role: IO, SHO, or ADMIN."
)

doc.add_heading("3.2 Browsing FIRs", level=2)
doc.add_paragraph(
    'Below the upload zone is the "FIR Browse" table. It shows all ingested FIRs '
    "in reverse chronological order with pagination (10 per page)."
)
doc.add_paragraph("Available columns:")
cols = [
    "FIR Number — the unique identifier extracted from the PDF",
    "District — extracted or inferred district name",
    "Police Station — extracted station name",
    "Status — one of: pending, classified, reviewed, review_needed",
    "NLP Category — the AI-predicted crime category with a confidence bar",
    "Completeness — extraction completeness percentage",
    "Date — ingestion timestamp",
]
for c in cols:
    doc.add_paragraph(c, style="List Bullet")

doc.add_paragraph(
    'Use the "Filter by district" text input to narrow results. '
    'Click "View" on any row to open a slide-over panel with the full FIR details '
    "including narrative text and NLP classification metadata."
)
doc.add_paragraph(
    "IO and SHO users only see FIRs from their own district. DYSP, SP, ADMIN, "
    "and READONLY users see all districts."
)

doc.add_heading("3.3 NLP Classification & Mismatch Detection", level=2)
doc.add_paragraph(
    "Each FIR is automatically classified into one of 11 crime categories: "
    "murder, assault, rape/sexual offences, domestic violence, kidnapping, theft, "
    "dacoity/robbery, fraud, cybercrime, narcotics, or other."
)
doc.add_paragraph(
    "The classification uses a 4-tier cascade:"
)
tiers = [
    "Tier 1 — Section Map: deterministic mapping of IPC/BNS sections to crime categories (confidence 1.0)",
    "Tier 2 — Fine-tuned MuRIL: IndicBERT model trained on labelled FIR narratives",
    "Tier 3 — Zero-shot NLI: multilingual mDeBERTa for cases where no training data exists",
    "Tier 4 — Keyword Heuristics: fallback pattern matching",
]
for t in tiers:
    doc.add_paragraph(t, style="List Bullet")

doc.add_paragraph(
    "When the section-inferred category differs from the NLP-predicted category, "
    "the FIR is flagged as 'review_needed' with a mismatch alert. An amber warning "
    "banner appears in the FIR detail view."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 4. CHARGESHEET MODULE
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading("4. Chargesheet Module", level=1)
doc.add_paragraph(
    'Navigate to "Charge Sheet" in the left sidebar. This module handles chargesheet '
    "PDF ingestion, AI-powered validation, evidence analysis, and the full review workflow."
)

doc.add_heading("4.1 Uploading a Chargesheet PDF", level=2)
steps = [
    'Drag a PDF onto the upload zone labelled "Drag & drop a Charge-Sheet PDF here", or click to browse.',
    "The system extracts text and parses structured fields: accused persons, charge sections, evidence items, witnesses, IO name, court name, filing date, and FIR reference.",
    "If the FIR reference number matches an existing FIR in the database, the chargesheet is automatically linked.",
    "The parsed result card shows: Court, District, IO, Accused count, Charges, and FIR link status.",
    "Required role: SHO, DYSP, SP, or ADMIN. IO users will receive a 403 error.",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_heading("4.2 Browsing Chargesheets", level=2)
doc.add_paragraph(
    "The Charge-Sheet Browse table shows all ingested chargesheets with filters for "
    "district (text input) and status (dropdown: All / Pending / Parsed / Reviewed / Flagged)."
)
doc.add_paragraph("Table columns: Court | District | PS | IO | Status | Accused (count) | Charges | Date")
doc.add_paragraph(
    'Click "View" on any row to navigate to the three-panel review page at '
    "/dashboard/chargesheet/[id]."
)

doc.add_heading("4.3 Reviewing a Chargesheet", level=2)
doc.add_paragraph(
    "The chargesheet review page has a three-panel layout designed for thorough, AI-assisted review."
)

doc.add_paragraph("Left Panel — Document Viewer", style="Intense Quote")
doc.add_paragraph(
    "Displays all parsed chargesheet fields in collapsible cards. Click any card header "
    "to expand or collapse it:"
)
cards = [
    "Case Details: IO name, court name, filing date, FIR link status",
    "Accused: table with Name, Age, and Role columns",
    "Charges: IPC/BNS section badges with descriptions",
    "Evidence: table with Type, Description, and Status (collected/pending)",
    "Witnesses: name, role badge, and statement summary",
]
for c in cards:
    doc.add_paragraph(c, style="List Bullet")

doc.add_paragraph("Right Panel — AI Recommendations", style="Intense Quote")
doc.add_paragraph(
    "After clicking 'Start Review', this panel loads AI-generated recommendations "
    "from two sources: Legal Validation and Evidence Gap Detection."
)
doc.add_paragraph("Three tabs are available:")
tabs = [
    "Legal: Findings from the legal cross-reference validator (7 rules). Each finding shows severity, section, description, recommendation, and action buttons.",
    "Evidence: Findings from the evidence gap detector. Shows a coverage meter at the top, then gap cards with tier badges (Rule-based or AI-suggested), and a list of present evidence items.",
    "Summary: Shows action counters (Accepted / Modified / Dismissed), an overall assessment text area, a 'Flag for senior review' checkbox, and the 'Complete Review' button.",
]
for t in tabs:
    doc.add_paragraph(t, style="List Bullet")

doc.add_paragraph("Bottom Panel — Audit Trail", style="Intense Quote")
doc.add_paragraph(
    "A collapsible panel showing the chronological audit log. Each entry shows: "
    "action type (color-coded dot), user, hash (truncated), and timestamp. "
    "Buttons for 'Verify Chain' (SP+ only) and 'Export CSV' (DySP+ only) are available."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 5. LEGAL VALIDATION
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading("5. Legal Validation", level=1)
doc.add_paragraph(
    "The legal cross-reference validator checks chargesheet sections against the linked "
    "FIR and the IPC/BNS legal database. It automatically runs when a review is started."
)

doc.add_heading("5.1 Validation Rules", level=2)
doc.add_paragraph("Seven rules are checked:")
add_table(
    ["Rule", "Severity", "What It Checks"],
    [
        ["RULE 1 — Section Mismatch", "WARNING", "Sections in chargesheet but absent from FIR (new sections added without supplementary statement)"],
        ["RULE 2 — Dropped Sections", "ERROR", "Sections in FIR but missing from chargesheet (potential dropped charges)"],
        ["RULE 3 — Invalid Combinations", "ERROR", "Mutually exclusive sections charged together (e.g., 302 + 304 on same victim)"],
        ["RULE 4 — Missing Companions", "WARNING", "Primary section charged without standard companion section (e.g., 302 without 201)"],
        ["RULE 5 — Procedural Gaps", "CRITICAL", "Required procedural steps not reflected in evidence (e.g., 376 without 164 CrPC statement)"],
        ["RULE 6 — IPC/BNS Mismatch", "ERROR", "Wrong act for case date (IPC on post-July 2024 case, or BNS on pre-July 2024 case)"],
        ["RULE 7 — Evidence Sufficiency", "WARNING", "Mandatory evidence for charged section missing from evidence list"],
    ],
)
doc.add_paragraph(
    "Rules 1 and 2 only run when the chargesheet is linked to a FIR. "
    "Rules 3–7 always run (internal consistency checks)."
)

doc.add_heading("5.2 Section Lookup", level=2)
doc.add_paragraph(
    "On the chargesheet review page (Charges tab), click any IPC/BNS section number "
    "to open a lookup tooltip showing:"
)
items = [
    "Section title and crime category",
    "Cognizable / bailable status",
    "Maximum sentence",
    "BNS equivalent (for IPC sections) or IPC equivalent (for BNS sections)",
    "Mandatory evidence items",
    "Companion sections",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 6. EVIDENCE GAP DETECTION
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading("6. Evidence Gap Detection", level=1)
doc.add_paragraph(
    "The evidence gap detector uses a two-tier architecture to identify missing "
    "evidence categories. It runs automatically when a review is started."
)

doc.add_heading("6.1 Rule-Based Detection (Tier 1)", level=2)
doc.add_paragraph(
    "Based on the crime category and charged sections, the system looks up expected "
    "evidence from the taxonomy (20 categories). Each evidence item in the chargesheet "
    "is classified into a canonical category using keyword matching. "
    "Gaps = expected items minus present items."
)
doc.add_paragraph("Severity levels:")
sevs = [
    "Critical: mandatory evidence for the crime type (e.g., post-mortem for murder)",
    "Important: commonly expected evidence (e.g., CCTV footage for robbery)",
    "Supplementary: optional supporting evidence (e.g., confession statement)",
]
for s in sevs:
    doc.add_paragraph(s, style="List Bullet")

doc.add_heading("6.2 ML Pattern Detection (Tier 2)", level=2)
doc.add_paragraph(
    "A scikit-learn multi-label classifier (TF-IDF + LogisticRegression) trained on "
    "synthetic data predicts the probability of each evidence category being present. "
    "Low-probability items not already flagged by Tier 1 are surfaced as 'AI-suggested' "
    "gaps. These appear with a dashed border badge to distinguish them from rule-based findings."
)

doc.add_heading("6.3 Evidence Taxonomy", level=2)
doc.add_paragraph(
    "The system recognises 20 evidence categories:"
)
categories = [
    "Post-mortem report", "Scene of crime report", "Forensic report", "CCTV footage",
    "Witness statements (161 CrPC)", "Victim statement (164 CrPC)", "Medical examination",
    "Weapon recovery", "Electronic evidence", "Financial records", "Identification parade",
    "Call detail records", "Narcotics test report", "Property valuation", "Confession statement",
    "Site plan / map", "DNA report", "Fingerprint report", "Seizure memo", "FSL report",
]
for c in categories:
    doc.add_paragraph(c, style="List Bullet")

doc.add_paragraph(
    "The coverage meter on the Evidence tab shows the percentage of expected evidence "
    "that is present, colour-coded: green (≥80%), yellow (50–80%), red (<50%)."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 7. REVIEW WORKFLOW
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading("7. Review Workflow", level=1)
doc.add_paragraph(
    "The review workflow is a structured process for investigators to act on every "
    "AI recommendation before finalising a chargesheet."
)

doc.add_heading("7.1 Starting a Review", level=2)
steps = [
    "Navigate to the chargesheet detail page (/dashboard/chargesheet/[id]).",
    'Click the "Start Review" button in the header.',
    "The chargesheet status changes to 'under_review'.",
    "The system automatically runs legal validation and evidence gap detection.",
    "Recommendations appear in the right panel under the Legal and Evidence tabs.",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_heading("7.2 Acting on Recommendations", level=2)
doc.add_paragraph("Each recommendation card has three action buttons:")
add_table(
    ["Action", "What Happens", "Required Input"],
    [
        ["Accept", "The recommendation is accepted as-is. The card greys out.", "None"],
        ["Modify", "A text input appears. Enter your modified version.", "Modified text (required)"],
        ["Dismiss", "A text input appears. Enter your reason for dismissal.", "Reason (required)"],
    ],
)
doc.add_paragraph(
    "Each action is immediately recorded in the audit trail. You cannot act on "
    "the same recommendation twice (the system returns a 409 Conflict error)."
)

doc.add_heading("7.3 Completing a Review", level=2)
steps = [
    'Switch to the "Summary" tab in the right panel.',
    "Review the action counters (Accepted / Modified / Dismissed).",
    "Enter overall assessment notes in the text area (optional).",
    'Check "Flag for senior review" if the case needs escalation.',
    'Click "Complete Review".',
    "The chargesheet status updates to 'reviewed' (or 'flagged' if the checkbox was checked).",
    "A final audit entry is logged.",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 8. AUDIT TRAIL
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading("8. Audit Trail", level=1)
doc.add_paragraph(
    "Every action on a chargesheet is logged in a tamper-evident audit chain. "
    "Each entry is linked to the previous one via a SHA-256 hash, making it "
    "impossible to alter or delete entries without breaking the chain."
)

doc.add_heading("8.1 Viewing the Audit Log", level=2)
doc.add_paragraph(
    "On the chargesheet review page, click 'Show Audit Trail' at the bottom. "
    "The timeline shows each action with a colour-coded dot, the user who performed it, "
    "a truncated hash, and the timestamp."
)
doc.add_paragraph("Logged actions include:")
actions = [
    "REVIEW_STARTED — review session begun",
    "RECOMMENDATION_ACCEPTED — a recommendation was accepted",
    "RECOMMENDATION_MODIFIED — a recommendation was modified with custom text",
    "RECOMMENDATION_DISMISSED — a recommendation was dismissed with reason",
    "REVIEW_COMPLETED — review finalised with 'reviewed' status",
    "REVIEW_FLAGGED — review finalised with 'flagged' status for senior review",
    "EXPORT_GENERATED — audit trail exported as CSV",
]
for a in actions:
    doc.add_paragraph(a, style="List Bullet")
doc.add_paragraph("Required role to view audit: DySP, SP, or ADMIN.")

doc.add_heading("8.2 Verifying Chain Integrity", level=2)
doc.add_paragraph(
    'Click "Verify Chain" in the audit panel. The system walks every entry, '
    "recomputes each SHA-256 hash, and checks that each entry's previous_hash "
    "matches the prior entry's hash. A green badge shows 'Chain verified' if intact, "
    "or a red alert shows where the chain first breaks."
)
doc.add_paragraph("Required role: SP or ADMIN only.")

doc.add_heading("8.3 Exporting for Court", level=2)
doc.add_paragraph(
    'Click "Export CSV" to download the full audit chain as a CSV file. '
    "The export includes all fields: entry ID, chargesheet ID, user, action, "
    "details, IP address, user agent, previous hash, entry hash, and timestamp. "
    "This file can be submitted as part of court proceedings to demonstrate "
    "the integrity of the review process."
)
doc.add_paragraph("Required role: DySP, SP, or ADMIN.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 9. ROLES & PERMISSIONS
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading("9. Roles & Permissions", level=1)
doc.add_paragraph(
    "ATLAS uses Role-Based Access Control (RBAC) with six roles. "
    "Each role determines which features and data a user can access."
)
add_table(
    ["Role", "FIR Access", "CS Upload", "Review", "Audit View", "Chain Verify", "Admin"],
    [
        ["IO", "Own district", "No", "No", "No", "No", "No"],
        ["SHO", "Own district", "Yes", "Yes", "No", "No", "No"],
        ["DySP", "All districts", "Yes", "Yes", "Yes", "No", "No"],
        ["SP", "All districts", "Yes", "Yes", "Yes", "Yes", "No"],
        ["ADMIN", "All districts", "Yes", "Yes", "Yes", "Yes", "Yes"],
        ["READONLY", "All districts", "No", "No", "No", "No", "No"],
    ],
)
doc.add_paragraph(
    "IO and SHO users are district-scoped: they only see FIRs and chargesheets "
    "from their assigned district. DySP and above see all districts."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 10. SERVICE URLs & MONITORING
# ══════════════════════════════════════════════════════════════════════════

doc.add_heading("10. Service URLs & Monitoring", level=1)
doc.add_paragraph("After running 'docker compose up -d', the following services are available:")
add_table(
    ["Service", "URL", "Credentials"],
    [
        ["Frontend (Web App)", "http://localhost:3000", "See Section 1.3"],
        ["Backend API", "http://localhost:8000", "JWT token required"],
        ["API Documentation", "http://localhost:8000/docs", "Interactive Swagger UI"],
        ["Grafana (Monitoring)", "http://localhost:3001", "admin / atlasadmin"],
        ["Prometheus (Metrics)", "http://localhost:9090", "No auth required"],
        ["Label Studio (Annotation)", "http://localhost:8080", "atlas@atlas.local / atlasadmin"],
        ["MLflow (Experiments)", "http://localhost:5000", "No auth required"],
        ["PostgreSQL", "localhost:5433", "atlas / atlaspass"],
        ["Redis", "localhost:6380", "No auth"],
        ["MongoDB", "localhost:27017", "No auth"],
    ],
)
doc.add_paragraph(
    "Grafana is pre-configured with Prometheus as a data source. It tracks "
    "API request counts, latency percentiles, and error rates for the backend."
)
doc.add_paragraph(
    "MLflow tracks ML experiment metrics for the FIR classification model and "
    "the evidence gap detection model."
)

# ══════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════

out_path = Path(__file__).parent / "ATLAS_User_Guide.docx"
doc.save(str(out_path))
print(f"Saved: {out_path}")
print(f"Size: {out_path.stat().st_size / 1024:.0f} KB")
