#!/usr/bin/env python3
"""Generate the ATLAS Architecture & AI Decisions document as .docx."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path

doc = Document()

# ── Styles ────────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    doc.add_paragraph()

_DIAGRAMS_DIR = Path(__file__).parent / "_diagrams"

def add_diagram_image(image_filename, caption=None, width_inches=6.2):
    """Add a PNG diagram image with optional caption."""
    img_path = _DIAGRAMS_DIR / image_filename
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(width_inches))
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(caption)
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            r.italic = True
    else:
        doc.add_paragraph(f"[Diagram not found: {image_filename}]")
    doc.add_paragraph()

def add_decision_box(decision, rationale, alternatives):
    """Add a structured decision record."""
    p = doc.add_paragraph()
    r = p.add_run("DECISION: ")
    r.bold = True; r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    p.add_run(decision)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("RATIONALE: ")
    r2.bold = True; r2.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
    p2.add_run(rationale)
    p3 = doc.add_paragraph()
    r3 = p3.add_run("ALTERNATIVES CONSIDERED: ")
    r3.bold = True; r3.font.color.rgb = RGBColor(0xBF, 0x36, 0x0C)
    p3.add_run(alternatives)
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════
doc.add_paragraph(); doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("ATLAS Platform"); r.font.size = Pt(36); r.bold = True; r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Architecture, Control Flow & AI Design Decisions"); r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x4A, 0x6F, 0xA5)
doc.add_paragraph()
d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = d.add_run("A comprehensive technical reference documenting every\narchitectural choice, AI model selection, data flow, and\ndesign trade-off in the ATLAS criminal justice platform.")
r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
doc.add_paragraph(); doc.add_paragraph()
v = doc.add_paragraph(); v.alignment = WD_ALIGN_PARAGRAPH.CENTER
v.add_run("Version 2.0 — April 2026\nBased on implementation analysis (not documentation)").font.size = Pt(11)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc = [
    "1. System Overview & High-Level Architecture",
    "2. Infrastructure & Deployment Architecture",
    "    2.1 Why Docker Compose (Not Kubernetes)",
    "    2.2 Service Topology & Port Mapping",
    "    2.3 Database Architecture: The Polyglot Persistence Strategy",
    "    2.4 Health Checks & Dependency Ordering",
    "3. Backend Architecture",
    "    3.1 Why FastAPI (Not Django/Flask)",
    "    3.2 Connection Management: Singleton vs Pool",
    "    3.3 Raw SQL vs ORM: The psycopg2 Decision",
    "    3.4 JSONB vs Normalized Tables",
    "4. Frontend Architecture",
    "    4.1 Why Next.js 14 (Not React SPA / Angular)",
    "    4.2 Component Library: shadcn/ui + Tailwind",
    "5. Authentication & Security",
    "    5.1 JWT Architecture",
    "    5.2 RBAC: The 6-Role Permission Matrix",
    "    5.3 PII Masking: Role-Graduated Privacy",
    "    5.4 Victim Identity Protection (BNS Section 73)",
    "6. PDF Extraction Pipeline",
    "    6.1 The Three-Stage Fallback Chain",
    "    6.2 Why pdfplumber + Tesseract (Not AWS Textract)",
    "    6.3 OCR Configuration Decisions",
    "7. FIR Parsing Engine",
    "    7.1 Anchor-Based Regex vs Positional Extraction",
    "    7.2 Gujarati Numeral Handling",
    "    7.3 OCR Error Correction",
    "8. AI/ML Classification System",
    "    8.1 The 4-Tier Classification Cascade",
    "    8.2 Tier 1: Section Map (Deterministic)",
    "    8.3 Tier 2: Fine-Tuned MuRIL (Why Not mBERT / XLM-R)",
    "    8.4 Tier 3: Zero-Shot NLI (Why mDeBERTa)",
    "    8.5 Tier 4: Keyword Heuristics (Safety Net)",
    "    8.6 Confidence Thresholds & Fallback Logic",
    "9. NLP Preprocessing Pipeline",
    "    9.1 Language Detection (fastText lid.176.bin)",
    "    9.2 Transliteration (IndicXlit)",
    "    9.3 Sentence Splitting (Indic NLP Library)",
    "10. Chargesheet Processing",
    "    10.1 Parser Architecture",
    "    10.2 Auto-Linkage to FIRs",
    "11. Legal Cross-Reference Validation Engine",
    "    11.1 The 7-Rule Validation Framework",
    "    11.2 Legal Database (sections.json): 73 Sections",
    "    11.3 IPC-to-BNS Transition Handling",
    "12. Evidence Gap Detection System",
    "    12.1 Two-Tier Architecture",
    "    12.2 Evidence Taxonomy: 20 Categories",
    "    12.3 ML Model: Why Logistic Regression (Not Deep Learning)",
    "    12.4 Training Pipeline & Synthetic Data",
    "13. Review Workflow & Audit Trail",
    "    13.1 The SHA-256 Hash Chain",
    "    13.2 Why Append-Only (Not Mutable Logs)",
    "    13.3 Accept / Modify / Dismiss Pattern",
    "    13.4 Court-Admissible Export",
    "14. Observability Stack",
    "    14.1 Prometheus + Grafana",
    "    14.2 MLflow Experiment Tracking",
    "    14.3 Structured Logging (structlog)",
    "15. Control Flow Diagrams",
    "    15.1 FIR Ingestion End-to-End",
    "    15.2 Chargesheet Review End-to-End",
    "    15.3 Classification Cascade Flow",
    "16. Comprehensive FAQ",
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(1)
    if not item.startswith("    "): p.runs[0].bold = True
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 1. SYSTEM OVERVIEW
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("1. System Overview & High-Level Architecture", level=1)
doc.add_paragraph(
    "ATLAS (Automated Tracking and Legal Analysis System) is an AI-assisted platform "
    "for processing First Information Reports (FIRs) and charge-sheets in the Indian "
    "criminal justice system. It is designed for on-premise deployment at district police "
    "hubs in Gujarat, operating entirely on CPU hardware (16 GB RAM, no GPU) with no "
    "external API dependencies."
)
doc.add_paragraph("The system has five major subsystems:")
subs = [
    ("Document Ingestion", "PDF extraction with OCR for scanned Gujarati/English police documents"),
    ("NLP Classification", "4-tier AI cascade for crime category classification"),
    ("Legal Validation", "7-rule engine cross-referencing chargesheet sections against FIR and IPC/BNS law"),
    ("Evidence Gap Detection", "Two-tier (rule + ML) system identifying missing evidence per crime type"),
    ("Review & Audit", "Investigator review workflow with SHA-256 hash-chain audit trail for court admissibility"),
]
for title, desc in subs:
    p = doc.add_paragraph()
    r = p.add_run(f"{title}: "); r.bold = True
    p.add_run(desc)

add_diagram_image("01_system_architecture.png", "Figure 1: ATLAS Platform System Architecture")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 2. INFRASTRUCTURE
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("2. Infrastructure & Deployment Architecture", level=1)

doc.add_heading("2.1 Why Docker Compose (Not Kubernetes)", level=2)
add_decision_box(
    "Use Docker Compose for orchestration, not Kubernetes.",
    "ATLAS is deployed at individual district police hubs — small-scale, single-node servers "
    "with 16 GB RAM. Kubernetes adds operational complexity (etcd, control plane, ingress "
    "controllers) that is unnecessary for a single-machine deployment. Docker Compose provides "
    "service dependency ordering via health checks, automatic restarts, and named volumes — "
    "all the primitives needed. A future migration to K8s is straightforward since each "
    "service already runs in its own container.",
    "Kubernetes (over-engineered for single-node), systemd services (no container isolation), "
    "Docker Swarm (deprecated ecosystem)."
)

doc.add_heading("2.2 Service Topology & Port Mapping", level=2)
add_table(
    ["Service", "Image", "Internal Port", "Host Port", "Purpose"],
    [
        ["backend", "python:3.11-slim (custom)", "8000", "8000", "FastAPI application server"],
        ["frontend", "node:20-alpine (custom)", "3000", "3000", "Next.js web application"],
        ["db", "postgres:15", "5432", "5433", "Primary relational database"],
        ["redis", "redis:7", "6379", "6380", "Session cache & rate limiting"],
        ["mongodb", "mongo:7", "27017", "27017", "Raw OCR text storage"],
        ["prometheus", "prom/prometheus:v2.50.0", "9090", "9090", "Metrics collection"],
        ["grafana", "grafana/grafana:10.3.0", "3000", "3001", "Metrics dashboards"],
        ["mlflow", "ghcr.io/mlflow/mlflow:v2.11.1", "5000", "5000", "ML experiment tracking"],
        ["labelstudio", "heartexlabs/label-studio", "8080", "8080", "Annotation tool"],
    ],
)
doc.add_paragraph(
    "Host ports 5433 and 6380 are remapped to avoid conflicts with other services "
    "on development machines. Internal Docker-network connections remain on standard ports."
)

doc.add_heading("2.3 Database Architecture: The Polyglot Persistence Strategy", level=2)
add_decision_box(
    "Use three different databases: PostgreSQL (structured), MongoDB (raw text), Redis (cache).",
    "Police documents have two fundamentally different data shapes: (1) structured fields "
    "(FIR number, sections, dates) that require ACID transactions, foreign keys, and SQL "
    "queries — ideal for PostgreSQL; (2) raw OCR text blobs (5-50 KB each) that need "
    "fast document-level writes with no joins — ideal for MongoDB. Redis provides sub-millisecond "
    "caching for JWT session state and frequently-accessed lookups. A single database could "
    "handle all three, but at the cost of either schema complexity (PostgreSQL for documents) "
    "or transactional integrity loss (MongoDB for everything).",
    "PostgreSQL-only (TEXT columns for raw OCR — viable but slower for bulk text queries), "
    "MongoDB-only (loses FK constraints and ACID for FIR-to-chargesheet links)."
)

doc.add_paragraph("PostgreSQL tables (10 tables total):")
add_table(
    ["Table", "Primary Key", "Key Design Choice"],
    [
        ["firs", "UUID", "29 columns; JSONB for stolen_property; TEXT[] for primary_sections"],
        ["chargesheets", "UUID", "JSONB for accused, charges, evidence, witnesses (semi-structured)"],
        ["users", "UUID", "bcrypt password_hash; role-based district scoping"],
        ["audit_log", "BIGSERIAL", "Legacy audit log (generic, FIR-era)"],
        ["audit_log_chargesheet", "UUID", "SHA-256 hash chain; entry_hash UNIQUE; append-only"],
        ["recommendation_actions", "UUID", "FK to audit_log_chargesheet for traceability"],
        ["validation_reports", "UUID", "JSONB for findings and summary"],
        ["evidence_gap_reports", "UUID", "JSONB for gaps and present evidence"],
        ["complainants / accused", "UUID", "Normalized child tables of firs"],
        ["ocr_jobs", "UUID", "Job tracking for async processing"],
    ],
)

doc.add_heading("2.4 Health Checks & Dependency Ordering", level=2)
doc.add_paragraph(
    "Docker Compose uses health checks to enforce service startup order. The backend "
    "will not start until PostgreSQL passes pg_isready, Redis passes redis-cli ping, "
    "and MongoDB passes mongosh ping. This prevents connection errors on cold boot. "
    "Each service uses restart: unless-stopped for crash resilience."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 3. BACKEND ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("3. Backend Architecture", level=1)

doc.add_heading("3.1 Why FastAPI (Not Django / Flask)", level=2)
add_decision_box(
    "FastAPI 0.110.0 on Python 3.11 with Uvicorn ASGI server.",
    "FastAPI provides: (1) automatic OpenAPI/Swagger documentation at /docs — critical for "
    "a multi-team project; (2) native async support for I/O-bound OCR and ML inference; "
    "(3) Pydantic v2 integration for request/response validation with zero boilerplate; "
    "(4) dependency injection for RBAC via Depends(). Django was rejected because its ORM "
    "adds overhead we don't need (we use raw SQL), and its monolithic design conflicts with "
    "our modular file layout. Flask was rejected because it lacks native async support and "
    "has no built-in validation — we would need to add marshmallow/flask-restx manually.",
    "Django REST Framework (ORM overhead, synchronous by default), Flask + marshmallow "
    "(no async, manual validation), Express.js (would require rewriting ML pipeline in JS)."
)

doc.add_heading("3.2 Connection Management: Singleton vs Pool", level=2)
add_decision_box(
    "Use a lazy singleton psycopg2 connection, not a connection pool.",
    "The backend runs as a single Uvicorn worker inside Docker. A connection pool (e.g., "
    "psycopg2.pool.ThreadedConnectionPool) is designed for multi-process deployments where "
    "multiple workers compete for DB connections. In our single-worker container, a singleton "
    "connection with automatic reconnection on failure is simpler and equally performant. "
    "The session.py module detects stuck transactions (TRANSACTION_STATUS_INERROR) and rolls "
    "back automatically, preventing connection leaks.",
    "SQLAlchemy engine with pool (adds 3,000 lines of ORM abstraction for ~10 queries), "
    "psycopg2 ThreadedConnectionPool (unnecessary for single-worker deployment)."
)

doc.add_heading("3.3 Raw SQL vs ORM: The psycopg2 Decision", level=2)
add_decision_box(
    "Use raw parameterised SQL via psycopg2, not an ORM.",
    "ATLAS has ~10 tables with straightforward CRUD operations. An ORM like SQLAlchemy "
    "adds a migration framework, model declarations, relationship mapping, and session "
    "management — all for queries that are simple INSERT/SELECT statements. Raw psycopg2 "
    "with RealDictCursor gives us dict-based result rows with zero abstraction cost. "
    "All queries use parameterised placeholders (%s) to prevent SQL injection. "
    "Alembic is used only for migration versioning, not for ORM model tracking.",
    "SQLAlchemy ORM (too heavy for our query complexity), Django ORM (requires Django), "
    "Tortoise ORM (less mature ecosystem)."
)

doc.add_heading("3.4 JSONB vs Normalized Tables", level=2)
doc.add_paragraph(
    "Chargesheet data (accused persons, charges, evidence items, witnesses) is stored as "
    "JSONB columns rather than normalized child tables. This is a deliberate trade-off:"
)
add_table(
    ["Factor", "JSONB", "Normalized"],
    [
        ["Query complexity", "Single SELECT returns all data", "4 JOINs required"],
        ["Schema flexibility", "No migration for new fields", "ALTER TABLE for each field"],
        ["Write performance", "Single INSERT", "5 INSERTs (parent + 4 children)"],
        ["Referential integrity", "None (application-level)", "FK constraints"],
        ["Analytics queries", "jsonb_array_elements required", "Standard SQL aggregates"],
    ],
)
doc.add_paragraph(
    "JSONB was chosen because chargesheet sub-entities (accused, evidence) are always "
    "read and written as a complete unit — never queried independently. The FIR table, "
    "by contrast, uses normalized child tables (complainants, accused) because FIR data "
    "was designed first when the schema was simpler."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 4. FRONTEND
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("4. Frontend Architecture", level=1)

doc.add_heading("4.1 Why Next.js 14 (Not React SPA / Angular)", level=2)
add_decision_box(
    "Next.js 14 with App Router, TypeScript, and Tailwind CSS.",
    "Next.js provides server-side rendering (SSR) for the login page (SEO-irrelevant but "
    "improves initial load time on slow police network connections), file-based routing "
    "(each page is a directory — /dashboard/chargesheet/[id]/page.tsx), and built-in "
    "image optimization. The App Router enables layouts that persist the sidebar across "
    "page navigations without re-rendering. TypeScript catches type errors at build time "
    "rather than runtime — critical when 6 different interfaces (ChargeSheet, FIR, Accused, "
    "Evidence, Witness, ValidationFinding) are passed between components.",
    "Create React App (no SSR, no file routing), Angular (steeper learning curve, heavier "
    "bundle), Vue/Nuxt (smaller ecosystem for complex data tables)."
)

doc.add_heading("4.2 Component Library: shadcn/ui + Tailwind", level=2)
add_decision_box(
    "Use shadcn/ui (copy-paste components) with Tailwind CSS, not Material UI or Ant Design.",
    "shadcn/ui copies component source code into the project (components/ui/card.tsx, "
    "button.tsx, badge.tsx, etc.) rather than importing from node_modules. This means: "
    "(1) zero runtime dependency — the component library cannot break on upgrade; (2) full "
    "customization — we modified Card and Badge styles for police-specific status colours "
    "(review_needed amber, classified green, flagged red) without fighting library defaults; "
    "(3) small bundle — only the 6 components we use are included, not the full 50+ library.",
    "Material UI (1.2 MB bundle, opinionated styling), Ant Design (CJK-focused, heavy), "
    "Chakra UI (viable alternative, slightly larger)."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 5. AUTH & SECURITY
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("5. Authentication & Security", level=1)

doc.add_heading("5.1 JWT Architecture", level=2)
doc.add_paragraph(
    "Authentication uses JSON Web Tokens (HS256) with bcrypt password hashing. The system "
    "issues two tokens: an access token (30-minute expiry) and a refresh token (7-day expiry). "
    "The JWT payload contains: username (sub), role, district, and full_name — this avoids "
    "a database lookup on every API call."
)
add_decision_box(
    "HS256 symmetric signing, not RS256 asymmetric.",
    "HS256 uses a single shared secret for both signing and verification. This is appropriate "
    "because the backend is the only service that creates and validates tokens — there is no "
    "microservice architecture where separate services need to verify tokens independently. "
    "HS256 is faster than RS256 and requires no certificate management.",
    "RS256 (unnecessary key management complexity), OAuth2/OIDC (no external identity provider "
    "exists in police districts), session cookies (incompatible with stateless API design)."
)

doc.add_heading("5.2 RBAC: The 6-Role Permission Matrix", level=2)
add_table(
    ["Role", "Scope", "FIR", "CS Upload", "Review", "Validate", "Audit", "Chain Verify"],
    [
        ["IO", "Own district", "Read/Write", "No", "No", "No", "No", "No"],
        ["SHO", "Own district", "Read/Write", "Yes", "Yes", "Yes", "No", "No"],
        ["DySP", "All districts", "Read", "Yes", "Yes", "Yes", "Yes", "No"],
        ["SP", "All districts", "Read", "Yes", "Yes", "Yes", "Yes", "Yes"],
        ["ADMIN", "All districts", "Full", "Yes", "Yes", "Yes", "Yes", "Yes"],
        ["READONLY", "All districts", "Read", "No", "No", "No", "No", "No"],
    ],
)
doc.add_paragraph(
    "IO and SHO users are district-scoped: SQL queries automatically append "
    "'WHERE district = %(d)s' for these roles. This is enforced at the CRUD layer, "
    "not the API layer, so it cannot be bypassed by crafting direct API requests."
)

doc.add_heading("5.3 PII Masking: Role-Graduated Privacy", level=2)
doc.add_paragraph(
    "PII masking is applied to every FIR response before it leaves the API. The masking "
    "level depends on the caller's role. Aadhaar numbers are masked to [AADHAAR] for all "
    "non-SP/ADMIN roles. Phone numbers are partially masked to [PHONE-XXXX] showing only "
    "the last 4 digits. Complainant names are abbreviated to 'Firstname L.' for DySP and "
    "READONLY roles."
)

doc.add_heading("5.4 Victim Identity Protection (BNS Section 73)", level=2)
doc.add_paragraph(
    "For FIRs involving sexual offences (BNS sections 63-99 or IPC sections 376, 354 family), "
    "the complainant_name is unconditionally replaced with [VICTIM-PROTECTED] and "
    "place_address with [ADDRESS-PROTECTED] for ALL roles including SP and ADMIN. "
    "This complies with Section 228A CrPC / BNS Section 73, which prohibits disclosure of "
    "the victim's identity in sexual offence cases. The check runs before role-based masking, "
    "ensuring it cannot be bypassed."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 6. PDF EXTRACTION
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("6. PDF Extraction Pipeline", level=1)

doc.add_heading("6.1 The Three-Stage Fallback Chain", level=2)
add_diagram_image("03_pdf_extraction.png", "Figure 3: PDF Text Extraction — Three-Stage Fallback Chain")

doc.add_heading("6.2 Why pdfplumber + Tesseract (Not AWS Textract)", level=2)
add_decision_box(
    "Use pdfplumber for text-layer PDFs and Tesseract OCR for scanned PDFs. No cloud APIs.",
    "ATLAS runs on-premise in police districts with unreliable internet. Cloud OCR services "
    "(AWS Textract, Google Vision) require network access and create data sovereignty issues — "
    "FIR documents contain sensitive PII that cannot leave the district network. Tesseract "
    "provides adequate accuracy for the structured eGujCop FIR format and supports both "
    "English (eng) and Gujarati (guj) language packs. pdfplumber is tried first because "
    "digital PDFs from eGujCop have embedded text layers, making OCR unnecessary in ~60% of cases.",
    "AWS Textract (requires internet, data sovereignty violation), Google Cloud Vision (same), "
    "EasyOCR (PyTorch-based — too heavy for CPU deployment alongside the NLP models)."
)

doc.add_heading("6.3 OCR Configuration Decisions", level=2)
add_table(
    ["Parameter", "Value", "Why"],
    [
        ["Primary DPI", "300", "Optimal for Tesseract accuracy on printed text; higher values slow processing without improving accuracy"],
        ["Fallback DPI", "200", "PyMuPDF fallback uses lower DPI to reduce memory usage"],
        ["Languages", "eng+guj", "eGujCop forms mix English headers with Gujarati body text"],
        ["PSM mode", "--psm 6", "Assumes a single uniform block of text (fits columnar FIR layout)"],
        ["Min text threshold", "50 chars", "Below this, pdfplumber output is likely garbage from an image PDF"],
    ],
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 7. FIR PARSING
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("7. FIR Parsing Engine", level=1)

doc.add_heading("7.1 Anchor-Based Regex vs Positional Extraction", level=2)
add_decision_box(
    "Use anchor-based regex extraction (look for labels like 'District:', 'FIR No.') "
    "rather than positional extraction (character offsets / coordinates).",
    "eGujCop FIR PDFs have inconsistent layouts: different districts use different templates, "
    "OCR word-splitting varies by scan quality, and Gujarati text breaks differently than "
    "English. Positional extraction would require per-template coordinate mapping that "
    "breaks with any layout change. Anchor-based regex finds the label 'District' then "
    "captures the value after it, regardless of where it appears on the page. The parser "
    "has multiple fallback patterns per field (e.g., 'District' in English, 'જીલ્લો' in "
    "Gujarati, plus OCR-corrupted variants like 'Distric t').",
    "Template-based coordinate extraction (brittle to layout changes), ML-based NER for "
    "field extraction (requires training data we don't have for Gujarati FIR forms)."
)

doc.add_heading("7.2 Gujarati Numeral Handling", level=2)
doc.add_paragraph(
    "eGujCop FIRs use Gujarati numerals (૦-૯) for FIR numbers, dates, and monetary values. "
    "The parser converts these to ASCII digits before any pattern matching using Python's "
    "str.maketrans() with a 10-character mapping table. This runs in O(n) time and avoids "
    "the need for Gujarati-aware regex patterns."
)

doc.add_heading("7.3 OCR Error Correction", level=2)
doc.add_paragraph(
    "Tesseract frequently splits common English words in eGujCop headers (e.g., 'Distric t', "
    "'Polic e', 'Sta tion', 'Complai nant'). The parser applies 11 regex-based corrections "
    "before field extraction. These corrections are applied in _normalise() and match "
    "case-insensitively. The system also collapses horizontal whitespace (tabs/spaces to "
    "single space) while preserving newlines as field separators."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 8. AI/ML CLASSIFICATION
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("8. AI/ML Classification System", level=1)
doc.add_paragraph(
    "FIR classification uses a 4-tier cascade that balances accuracy, speed, and availability. "
    "Each tier is tried in order; the first tier that produces a confident result is used."
)

doc.add_heading("8.1 The 4-Tier Classification Cascade", level=2)
add_diagram_image("02_classification_cascade.png", "Figure 2: 4-Tier Classification Cascade with Confidence Thresholds")

doc.add_heading("8.2 Tier 1: Section Map (Deterministic)", level=2)
doc.add_paragraph(
    "When IPC/BNS sections are extracted from the FIR form, section_map.py maps them "
    "directly to crime categories. This is the most reliable method (confidence 1.0) because "
    "sections are legally defined. The mapping covers ~80 IPC sections and ~50 BNS sections. "
    "When multiple sections map to different categories, a priority hierarchy resolves ties: "
    "murder > rape_sexoff > dacoity_robbery > kidnapping > domestic_violence > assault > "
    "fraud > narcotics > cybercrime > theft."
)

doc.add_heading("8.3 Tier 2: Fine-Tuned MuRIL (Why Not mBERT / XLM-R)", level=2)
add_decision_box(
    "Fine-tune google/muril-base-cased (237M parameters) for Indic language classification.",
    "MuRIL (Multilingual Representations for Indian Languages) was pre-trained on 17 Indian "
    "languages including Gujarati and Hindi, plus English. This gives it superior Gujarati "
    "understanding compared to mBERT (104 languages, Gujarati under-represented) or XLM-R "
    "(100 languages, similar under-representation). MuRIL achieves 4-8% higher F1 on Indic "
    "NER benchmarks. The model is fine-tuned on labelled FIR narratives using HuggingFace "
    "Trainer with MLflow autologging. Class-weighted loss handles the imbalanced category "
    "distribution. The confidence threshold of 0.25 (vs random chance at 1/11 = 0.091) "
    "ensures only meaningful predictions are accepted.",
    "mBERT (Gujarati under-represented in pre-training), XLM-RoBERTa (same issue), "
    "IndicBERT-v1 (gated model, access restricted), Sarvam/OpenHathi (too large for "
    "16 GB CPU deployment)."
)

doc.add_heading("8.4 Tier 3: Zero-Shot NLI (Why mDeBERTa)", level=2)
add_decision_box(
    "Use MoritzLaurer/mDeBERTa-v3-base-mnli-xnli (~270 MB) for zero-shot classification.",
    "Zero-shot NLI requires no training data — it scores each candidate label as a "
    "natural-language hypothesis ('this is a case of murder (હત્યા)') against the FIR text. "
    "mDeBERTa was chosen because: (1) it supports 100+ languages including Gujarati via "
    "XNLI training; (2) it's 270 MB (fits in 16 GB RAM alongside MuRIL); (3) bilingual "
    "hypothesis strings (English + Gujarati) improve cross-lingual transfer. Input is "
    "truncated to 600 characters because: each of the 11 labels requires a separate forward "
    "pass, and the FIR header + opening facts in the first 600 chars capture enough signal. "
    "Inference runs in asyncio.run_in_executor() to avoid blocking the uvicorn event loop "
    "(10-30 seconds on CPU per classification).",
    "BART-large-mnli (English-only, misses Gujarati), smaller DistilBERT variants "
    "(poor multilingual coverage), calling an LLM API (violates on-premise constraint)."
)

doc.add_heading("8.5 Tier 4: Keyword Heuristics (Safety Net)", level=2)
doc.add_paragraph(
    "The final fallback uses simple substring matching against 10 keyword lists "
    "(7+ keywords per category, in both English and Gujarati). This always produces a result — "
    "it never returns None. Confidence is calculated as min(0.5 + match_count * 0.1, 0.85). "
    "This ensures every FIR gets classified even if no models are loaded."
)

doc.add_heading("8.6 Confidence Thresholds & Fallback Logic", level=2)
add_table(
    ["Tier", "Threshold", "Rationale"],
    [
        ["Section Map", "None (always 1.0)", "Sections are legal ground truth"],
        ["MuRIL", "≥ 0.25", "1/11 random = 0.091; 0.25 = ~2.7x random chance"],
        ["Zero-Shot NLI", "≥ 0.20", "NLI scores are naturally lower; 0.20 filters noise"],
        ["Heuristics", "None (always runs)", "Last resort — any result better than nothing"],
    ],
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 9. NLP PREPROCESSING
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("9. NLP Preprocessing Pipeline", level=1)

doc.add_heading("9.1 Language Detection (fastText lid.176.bin)", level=2)
add_decision_box(
    "Use fastText lid.176.bin (126 MB) for language detection, not langdetect or polyglot.",
    "fastText's language identification model covers 176 languages with single-line detection. "
    "It correctly identifies Gujarati (gu) vs Hindi (hi) vs English (en) in the mixed-script "
    "eGujCop FIR text. Python's langdetect library uses a probabilistic algorithm that "
    "returns inconsistent results for short Gujarati text. The model is loaded once as a lazy "
    "singleton and cached for the process lifetime.",
    "langdetect (inconsistent on short Gujarati text), polyglot (requires system-level ICU "
    "libraries, painful in Docker), cld3 (deprecated by Google)."
)

doc.add_heading("9.2 Transliteration (IndicXlit)", level=2)
doc.add_paragraph(
    "Some FIR narratives contain Romanised Gujarati (e.g., 'hatya' for 'હત્યા'). IndicXlit "
    "converts these to native Gujarati script before NLP processing, improving classification "
    "accuracy for MuRIL which was trained on native-script text."
)

doc.add_heading("9.3 Sentence Splitting (Indic NLP Library)", level=2)
doc.add_paragraph(
    "The Indic NLP Library provides sentence splitting that correctly handles Gujarati "
    "sentence boundaries (Devanagari danda '।', double danda '॥', and Gujarati purna viram '.'). "
    "Standard Python splitlines() or NLTK sent_tokenize would mis-split Gujarati text."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 10. CHARGESHEET PROCESSING
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("10. Chargesheet Processing", level=1)

doc.add_heading("10.1 Parser Architecture", level=2)
doc.add_paragraph(
    "The chargesheet parser (chargesheet_parser.py) mirrors the FIR parser's design: "
    "anchor-based regex extraction, Gujarati numeral conversion, OCR error correction, "
    "and best-effort extraction that never raises exceptions. It extracts: FIR reference "
    "number, court name, filing date, investigation officer, accused persons (name, age, "
    "address, role), IPC/BNS charge sections, evidence items (type, description, status), "
    "and witness schedule (name, role, statement summary)."
)

doc.add_heading("10.2 Auto-Linkage to FIRs", level=2)
doc.add_paragraph(
    "When the parser extracts a FIR reference number, the ingest endpoint automatically "
    "queries the firs table to find a matching record. If found, the chargesheet's fir_id "
    "foreign key is set, enabling cross-reference validation between the two documents. "
    "This linkage is critical for Rules 1 and 2 of the legal validator."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 11. LEGAL VALIDATION
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("11. Legal Cross-Reference Validation Engine", level=1)

doc.add_heading("11.1 The 7-Rule Validation Framework", level=2)
doc.add_paragraph(
    "The legal validator runs 7 rules that check a chargesheet against its linked FIR "
    "and the legal database. Rules 1-2 require a linked FIR; rules 3-7 are internal "
    "consistency checks that always run."
)
add_table(
    ["Rule", "Name", "Severity", "What It Catches"],
    [
        ["1", "Section Mismatch", "WARNING", "New sections added without supplementary statement (173(8) CrPC)"],
        ["2", "Dropped Sections", "ERROR", "FIR sections removed from chargesheet without B-report"],
        ["3", "Invalid Combinations", "ERROR", "Mutually exclusive sections (e.g., 302+304 on same victim)"],
        ["4", "Missing Companions", "WARNING", "Primary section without standard companion (e.g., 302 without 201)"],
        ["5", "Procedural Gaps", "CRITICAL", "Required evidence absent (e.g., 376 without 164 CrPC statement)"],
        ["6", "IPC/BNS Mismatch", "ERROR", "Wrong act for case date (IPC post-July 2024 or BNS pre-July 2024)"],
        ["7", "Evidence Sufficiency", "WARNING", "Mandatory evidence per section missing from evidence list"],
    ],
)

doc.add_heading("11.2 Legal Database (sections.json): 73 Sections", level=2)
doc.add_paragraph(
    "The legal database contains 73 IPC/BNS section entries and 7 CrPC/BNSS procedural "
    "sections. Each entry stores: IPC section, BNS equivalent, title, crime category, "
    "mandatory evidence list, companion sections, procedural requirements, cognizable/bailable "
    "status, sentencing range, and mutually exclusive sections. The database is loaded once "
    "at module import and indexed by both IPC and BNS section number for O(1) lookup."
)

doc.add_heading("11.3 IPC-to-BNS Transition Handling", level=2)
doc.add_paragraph(
    "India transitioned from IPC to BNS on July 1, 2024. Rule 6 of the validator checks "
    "the FIR registration date: if after July 1, 2024, chargesheet sections should use BNS "
    "numbering; if before, IPC. This catches a common prosecutor error during the transition "
    "period where officers mix old and new section numbers."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 12. EVIDENCE GAP DETECTION
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("12. Evidence Gap Detection System", level=1)

doc.add_heading("12.1 Two-Tier Architecture", level=2)
add_diagram_image("04_evidence_gap.png", "Figure 4: Two-Tier Evidence Gap Detection Architecture")

doc.add_heading("12.2 Evidence Taxonomy: 20 Categories", level=2)
doc.add_paragraph(
    "The taxonomy defines 20 canonical evidence categories (post_mortem_report, "
    "scene_of_crime_report, forensic_report, cctv_footage, etc.) with crime-type "
    "applicability and weight (critical / important / supplementary). Free-text evidence "
    "descriptions from chargesheets are mapped to these categories using keyword matching "
    "(37 English + Gujarati + Hindi keyword groups) with rapidfuzz fuzzy matching as fallback "
    "(threshold 75%)."
)

doc.add_heading("12.3 ML Model: Why Logistic Regression (Not Deep Learning)", level=2)
add_decision_box(
    "Use scikit-learn MultiOutputClassifier(LogisticRegression) with TF-IDF features.",
    "The evidence gap model must: (1) run on CPU in <1 second; (2) be <10 MB on disk; "
    "(3) be interpretable (prosecutors need to trust the suggestions). Logistic regression "
    "with TF-IDF satisfies all three. A deep learning model (e.g., fine-tuned BERT for "
    "multi-label classification) would be 400+ MB, take 5+ seconds on CPU, and provide "
    "no interpretability benefit for this task. The model achieves macro-F1 ≥ 0.65 on "
    "synthetic data, which is adequate for a 'suggestion' tier that supplements the "
    "rule-based gaps. TF-IDF with 3,000 features and (1,2)-ngrams captures both single "
    "keywords and two-word phrases.",
    "BERT multi-label (too large for CPU), XGBoost (marginally better F1 but opaque), "
    "random forest (comparable F1 but larger model file)."
)

doc.add_heading("12.4 Training Pipeline & Synthetic Data", level=2)
doc.add_paragraph(
    "The model is trained on 2,200 synthetic charge-sheet samples (200 per crime category "
    "x 11 categories). Each sample has: text features (crime + sections + narrative), and "
    "20 binary labels (one per evidence category). Evidence presence is generated with "
    "realistic probabilities: 'strong' cases (90% of expected evidence present), 'medium' "
    "(65%), and 'weak' (30%). Non-expected evidence appears with 5-15% probability to "
    "represent real-world data noise. The model is trained using 80/20 train/test split "
    "with metrics logged to MLflow."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 13. REVIEW & AUDIT
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("13. Review Workflow & Audit Trail", level=1)

doc.add_heading("13.1 The SHA-256 Hash Chain", level=2)
add_diagram_image("05_audit_chain.png", "Figure 5: SHA-256 Tamper-Evident Audit Hash Chain")

add_decision_box(
    "Use SHA-256 hash chain for audit trail, not a blockchain or database triggers.",
    "A full blockchain (consensus protocol, mining/staking) is overkill for a single-node "
    "append-only log. Database triggers could prevent row updates but can be disabled by a "
    "DBA. The SHA-256 chain provides: (1) tamper evidence — altering any entry breaks the "
    "chain from that point forward; (2) verifiability — the chain can be walked and verified "
    "with a single O(n) scan; (3) court admissibility — exported CSV includes all hashes, "
    "allowing independent verification. The GENESIS sentinel for the first entry avoids "
    "NULL-handling complexity.",
    "PostgreSQL triggers (bypassable by admin), full blockchain (unnecessary consensus "
    "overhead), Merkle tree (more complex, same tamper evidence for sequential logs)."
)

doc.add_heading("13.2 Why Append-Only (Not Mutable Logs)", level=2)
doc.add_paragraph(
    "The audit_log_chargesheet table has a UNIQUE constraint on entry_hash and is designed "
    "as append-only. The application code never issues UPDATE or DELETE on this table. "
    "Any modification to a past entry would change its hash, which would cause a mismatch "
    "in the next entry's previous_hash, which the verify_chain() method detects."
)

doc.add_heading("13.3 Accept / Modify / Dismiss Pattern", level=2)
doc.add_paragraph(
    "Every AI recommendation (from legal validation or evidence gap detection) must be "
    "explicitly actioned by the reviewer. Three actions are available:"
)
items = [
    ("Accept", "The recommendation is valid as-is. No additional input required."),
    ("Modify", "The recommendation is partially correct. Reviewer provides modified text explaining the actual finding. The modified_text field is required."),
    ("Dismiss", "The recommendation is not applicable. Reviewer must provide a reason for dismissal. This creates an audit record explaining why the AI suggestion was rejected."),
]
for title, desc in items:
    p = doc.add_paragraph()
    r = p.add_run(f"{title}: "); r.bold = True
    p.add_run(desc)
doc.add_paragraph(
    "Duplicate actions on the same recommendation are rejected with HTTP 409 Conflict. "
    "This prevents accidental double-submissions."
)

doc.add_heading("13.4 Court-Admissible Export", level=2)
doc.add_paragraph(
    "The export endpoint generates a CSV file containing every audit entry with full "
    "hash chain data. This file can be independently verified by any party with a SHA-256 "
    "implementation by recomputing each hash and checking the chain. The export itself "
    "is logged as an EXPORT_GENERATED audit entry, creating a record of who exported "
    "the trail and when."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 14. OBSERVABILITY
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("14. Observability Stack", level=1)

doc.add_heading("14.1 Prometheus + Grafana", level=2)
doc.add_paragraph(
    "The prometheus-fastapi-instrumentator middleware exposes /metrics on the backend, "
    "recording HTTP request counts, latency percentiles, and error rates per endpoint. "
    "Prometheus scrapes these metrics every 15 seconds. Grafana (port 3001) provides "
    "pre-configured dashboards for API health monitoring."
)

doc.add_heading("14.2 MLflow Experiment Tracking", level=2)
doc.add_paragraph(
    "MLflow (port 5000) tracks two experiment types: (1) FIR classification model training "
    "(MuRIL fine-tuning with per-epoch loss, validation F1, confusion matrix artifacts); "
    "(2) evidence gap model training (macro-F1, per-category precision/recall). Each "
    "inference call can optionally log prediction confidence and category to MLflow for "
    "ongoing model monitoring."
)

doc.add_heading("14.3 Structured Logging (structlog)", level=2)
doc.add_paragraph(
    "All backend logs use structlog with JSON output, ISO timestamps, and log level/name "
    "fields. This enables log aggregation and search by any field (e.g., filtering all "
    "logs for a specific chargesheet_id). Python's built-in logging is configured at INFO "
    "level to capture pipeline progress without excessive debug noise."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 15. CONTROL FLOW DIAGRAMS
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("15. Control Flow Diagrams", level=1)

doc.add_heading("15.1 FIR Ingestion End-to-End", level=2)
add_diagram_image("06_fir_ingestion_flow.png", "Figure 6: FIR Ingestion — End-to-End Control Flow")

doc.add_heading("15.2 Chargesheet Review End-to-End", level=2)
add_diagram_image("07_review_flow.png", "Figure 7: Chargesheet Review — End-to-End Control Flow")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 16. COMPREHENSIVE FAQ
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("16. Comprehensive FAQ", level=1)

faqs = [
    # Architecture
    ("Why does ATLAS use three different databases?",
     "PostgreSQL stores structured data with ACID transactions (FIRs, chargesheets, users). "
     "MongoDB stores raw OCR text blobs (5-50 KB each) that are written once and rarely read — "
     "a document store is optimal. Redis caches JWT session state for sub-millisecond lookups. "
     "Using one database for all three would either sacrifice transaction safety (MongoDB for "
     "everything) or performance (PostgreSQL TEXT columns for OCR blobs)."),

    ("Why not use an ORM like SQLAlchemy or Django?",
     "ATLAS has ~10 tables with straightforward INSERT/SELECT queries. An ORM would add "
     "thousands of lines of model declarations, migration tooling, and session management for "
     "queries that are simple parameterised SQL. Raw psycopg2 with RealDictCursor returns "
     "dict rows directly — no object mapping overhead. Alembic handles migrations without "
     "requiring the full ORM."),

    ("Why is the frontend a separate Next.js app instead of server-rendered FastAPI templates?",
     "Separation of concerns: the backend is a pure JSON API, and the frontend is a standalone "
     "React application. This means: (1) the API can be used by other clients (mobile app, CLI); "
     "(2) frontend developers can work independently; (3) the backend can be deployed/scaled "
     "separately from the frontend. Jinja2 templates would tightly couple UI logic to Python."),

    ("Why use JSONB columns for chargesheet data instead of normalized tables?",
     "Chargesheet sub-entities (accused, evidence, witnesses) are always read and written as "
     "a complete unit — never queried independently. JSONB avoids 4 JOINs per query and "
     "allows schema evolution without ALTER TABLE. FIR data uses normalized tables because "
     "it was designed earlier when the schema was simpler."),

    # AI / ML
    ("Why does ATLAS use a 4-tier classification cascade instead of a single model?",
     "No single model handles all scenarios: (1) when IPC/BNS sections are present, they are "
     "legally authoritative (confidence 1.0) — no ML needed; (2) when sections are absent, "
     "the fine-tuned MuRIL model handles Gujarati text best but requires training data; "
     "(3) when no training data exists for a category, zero-shot NLI classifies without any "
     "labelled examples; (4) when no models are loaded (cold start, failed download), keyword "
     "heuristics ensure every FIR still gets classified. Each tier gracefully degrades to the "
     "next."),

    ("Why MuRIL and not mBERT or XLM-RoBERTa?",
     "MuRIL (Multilingual Representations for Indian Languages) was specifically pre-trained "
     "on 17 Indian languages including Gujarati and Hindi. mBERT covers 104 languages but "
     "under-represents Gujarati due to its smaller pre-training corpus. MuRIL achieves 4-8% "
     "higher F1 on Indic NER benchmarks. Both have the same architecture (BERT-base, 237M "
     "params) so inference speed is identical."),

    ("Why is zero-shot NLI used instead of training on more data?",
     "Real labelled FIR data is scarce (requires domain experts to annotate Gujarati legal "
     "text). Zero-shot NLI requires zero training data — it scores 'is this text about murder?' "
     "as an entailment hypothesis. This provides immediate classification capability while "
     "the annotation pipeline (Label Studio) collects real labelled data for future MuRIL "
     "fine-tuning."),

    ("Why Logistic Regression for evidence gap detection instead of a neural network?",
     "The evidence gap model must: (1) run in <1 second on CPU; (2) be <10 MB on disk; "
     "(3) be interpretable. Logistic regression with TF-IDF meets all three. A BERT-based "
     "multi-label classifier would be 400+ MB, take 5+ seconds per inference, and offer no "
     "interpretability benefit. The model is a 'suggestion' tier supplementing deterministic "
     "rules — high precision is less critical than low latency."),

    ("What are the 11 crime categories and how were they chosen?",
     "The categories align with BNS (Bharatiya Nyaya Sanhita 2023) chapter structure: "
     "murder (Ch.VI), assault (Ch.XVI), theft (Ch.XVII), fraud (Ch.XVII), rape/sexual offences "
     "(Ch.V), kidnapping (Ch.XIV), dacoity/robbery (Ch.XVII), domestic violence (PWDVA 2005), "
     "cybercrime (IT Act), narcotics (NDPS Act), and 'other'. These cover >95% of FIR filings "
     "in Gujarat based on NCRB data."),

    ("How does the system handle Gujarati text?",
     "At five levels: (1) OCR uses Tesseract with the Gujarati language pack (eng+guj); "
     "(2) Gujarati numerals (૦-૯) are converted to ASCII before parsing; (3) fastText detects "
     "Gujarati vs Hindi vs English; (4) IndicXlit transliterates Romanised Gujarati to native "
     "script; (5) MuRIL and mDeBERTa both support Gujarati natively in their vocabulary. "
     "Keyword heuristics include Gujarati terms (e.g., 'હત્યા' for murder, 'ચોરી' for theft)."),

    ("What happens if all AI models fail to load?",
     "The system degrades gracefully. If MuRIL fails to load, zero-shot NLI is tried. If "
     "zero-shot fails, keyword heuristics always work (no model needed). If fastText fails, "
     "language detection returns 'unknown' but parsing continues. Evidence gap detection "
     "Tier 1 (rule-based) works without any ML model. The system never blocks ingestion "
     "due to an AI failure — the FIR/chargesheet is always stored with whatever data could "
     "be extracted."),

    # Security
    ("How is PII protected in the system?",
     "Three layers: (1) role-based masking — Aadhaar numbers are masked for non-SP/ADMIN "
     "roles, phone numbers show only last 4 digits; (2) victim identity protection — for "
     "sexual offence cases (BNS §63-99), complainant names are unconditionally replaced with "
     "[VICTIM-PROTECTED] for ALL roles including ADMIN, per BNS §73; (3) narrative masking — "
     "PII patterns are redacted in free-text fields before storage."),

    ("How does the audit trail prevent tampering?",
     "Each audit entry's hash includes the previous entry's hash (SHA-256 chain). Altering "
     "any entry changes its hash, which breaks the chain at the next entry. The verify_chain() "
     "method walks all entries and recomputes each hash to detect any break. The audit table "
     "uses a UNIQUE constraint on entry_hash, preventing duplicate or replacement entries. "
     "The exported CSV includes full hashes for independent verification."),

    ("Why HS256 JWT instead of RS256 or OAuth2?",
     "HS256 uses a single shared secret. Since the backend is the only service that creates "
     "and validates tokens (no microservice architecture), symmetric signing is simpler and "
     "faster than RS256. OAuth2/OIDC requires an external identity provider, which doesn't "
     "exist in district police networks."),

    # Operations
    ("What are the hardware requirements?",
     "Minimum: 16 GB RAM, 4-core CPU, 50 GB disk. The largest memory consumers are: "
     "MuRIL model (~500 MB), mDeBERTa model (~270 MB), and PostgreSQL (~500 MB with data). "
     "No GPU is required — all inference runs on CPU using PyTorch CPU-only wheels."),

    ("How do I reset the database?",
     "Run: docker compose down -v && docker compose up --build. The -v flag removes named "
     "volumes (postgres_data, mongodb_data). On restart, init_schema.sql recreates all "
     "tables and 02_seed_users.sql inserts the default users."),

    ("How do I add a new user?",
     "Insert into the users table with a bcrypt-hashed password. The hash for 'atlas2025' "
     "is in 02_seed_users.sql. Supported roles: IO, SHO, DYSP, SP, ADMIN, READONLY. "
     "District and police_station fields control data scoping for IO/SHO roles."),

    ("What is the IPC-to-BNS transition and how does ATLAS handle it?",
     "India replaced the Indian Penal Code (IPC, 1860) with the Bharatiya Nyaya Sanhita "
     "(BNS, 2023) on July 1, 2024. All section numbers changed (e.g., IPC 302 → BNS 103). "
     "ATLAS handles this with: (1) dual section maps covering both IPC and BNS numbering; "
     "(2) a 73-entry cross-reference database mapping each IPC section to its BNS equivalent; "
     "(3) Rule 6 in the legal validator which flags chargesheets using the wrong act for the "
     "case registration date."),

    ("How does the evidence gap detector differ from Rule 7 in the legal validator?",
     "Rule 7 checks mandatory evidence per charged section (deterministic, per-section). "
     "The evidence gap detector checks evidence per crime category (broader scope), uses "
     "a 20-category taxonomy (vs the legal DB's free-text mandatory_evidence lists), and "
     "includes an ML tier that suggests non-obvious evidence based on patterns in similar "
     "cases. They are complementary: Rule 7 catches section-specific gaps, the evidence "
     "detector catches category-level gaps and ML-suggested additions."),

    ("Can ATLAS work offline (no internet)?",
     "Yes, completely. All ML models are downloaded once and cached locally in backend/models/. "
     "The Docker containers run with TRANSFORMERS_OFFLINE=1 and HF_DATASETS_OFFLINE=1 "
     "environment variables. No API calls are made to external services during normal "
     "operation. The only internet requirement is the initial docker compose build "
     "(to pull images) and model download."),

    ("What is the expected accuracy of the NLP classification?",
     "Tier 1 (section map) is 100% accurate by definition. Tier 2 (MuRIL) achieves "
     "macro-F1 of ~0.41 on synthetic data (limited by training data size — expected to "
     "improve with real annotations). Tier 3 (zero-shot NLI) provides reasonable "
     "classification without any training. Tier 4 (heuristics) is a best-effort fallback. "
     "The cascade design means the most accurate method available is always used first."),

    ("How are the 7 validation rules triggered?",
     "All 7 rules run automatically when a reviewer clicks 'Start Review' on the chargesheet "
     "detail page. Rules 1-2 (section mismatch, dropped sections) only produce findings if "
     "the chargesheet is linked to a FIR via fir_id. Rules 3-7 always run as internal "
     "consistency checks. Each finding has a severity (CRITICAL / ERROR / WARNING) and a "
     "specific recommendation for the reviewer."),
]

for q, a in faqs:
    p = doc.add_paragraph()
    r = p.add_run(f"Q: {q}")
    r.bold = True
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    doc.add_paragraph(f"A: {a}")
    doc.add_paragraph()  # spacer

# ══════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════
out_path = Path(__file__).parent / "ATLAS_Architecture_and_AI_Decisions.docx"
doc.save(str(out_path))
print(f"Saved: {out_path}")
print(f"Size: {out_path.stat().st_size / 1024:.0f} KB")
